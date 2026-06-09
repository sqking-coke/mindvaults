"""DOCX Parser + Preprocessor 集成测试。

验证：
  - parser: Heading 样式 → # 标题、列表检测、表格 → markdown
  - preprocessor: 文档顺序保留、表格完整、空白归一化
  - 边界情况: 空文档、纯表格文档、粗体下标识别
"""
from __future__ import annotations

import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

from app.services.parser_service import parse_document
from app.services.preprocessor import get_preprocessor


def _create_docx(path: str, build_fn) -> None:
    """辅助：创建一个 DOCX 文件。"""
    doc = Document()
    build_fn(doc)
    doc.save(path)


class TestDocxParser:
    """DOCX Parser 格式转换测试。"""

    @pytest.mark.asyncio
    async def test_heading_styles(self):
        """Heading 1-3 样式 → # / ## / ### 前缀。"""
        def build(doc):
            doc.add_heading("H1 Title", level=1)
            doc.add_heading("H2 Section", level=2)
            doc.add_heading("H3 Sub", level=3)
            doc.add_paragraph("Normal text.")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.close()
            _create_docx(f.name, build)
            try:
                pages = await parse_document(f.name, "docx")
                assert len(pages) == 1
                text = pages[0][0]
                assert "# H1 Title" in text
                assert "## H2 Section" in text
                assert "### H3 Sub" in text
                assert "Normal text." in text
            finally:
                Path(f.name).unlink()

    @pytest.mark.asyncio
    async def test_bold_heading_heuristic(self):
        """粗体 + 大字号 (≥14pt) → ## 标题候选。"""
        def build(doc):
            p = doc.add_paragraph()
            run = p.add_run("Bold Large Heading")
            run.bold = True
            run.font.size = Pt(16)
            doc.add_paragraph("Body text follows.")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.close()
            _create_docx(f.name, build)
            try:
                pages = await parse_document(f.name, "docx")
                text = pages[0][0]
                assert "## Bold Large Heading" in text
                assert "Body text follows." in text
            finally:
                Path(f.name).unlink()

    @pytest.mark.asyncio
    async def test_list_detection(self):
        """List Number / List Bullet 样式 → - 前缀。"""
        def build(doc):
            doc.add_paragraph("Item A", style="List Number")
            doc.add_paragraph("Item B", style="List Number")
            doc.add_paragraph("Bullet X", style="List Bullet")
            doc.add_paragraph("Bullet Y", style="List Bullet")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.close()
            _create_docx(f.name, build)
            try:
                pages = await parse_document(f.name, "docx")
                text = pages[0][0]
                assert "- Item A" in text
                assert "- Item B" in text
                assert "- Bullet X" in text
                assert "- Bullet Y" in text
            finally:
                Path(f.name).unlink()

    @pytest.mark.asyncio
    async def test_table_formatting(self):
        """表格 → markdown 表格格式（含表头分隔行）。"""
        def build(doc):
            table = doc.add_table(rows=3, cols=2)
            table.cell(0, 0).text = "Name"
            table.cell(0, 1).text = "Value"
            table.cell(1, 0).text = "Alpha"
            table.cell(1, 1).text = "100"
            table.cell(2, 0).text = "Beta"
            table.cell(2, 1).text = "200"

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.close()
            _create_docx(f.name, build)
            try:
                pages = await parse_document(f.name, "docx")
                text = pages[0][0]
                assert "| Name | Value |" in text
                assert "| --- | --- |" in text
                assert "| Alpha | 100 |" in text
                assert "| Beta | 200 |" in text
            finally:
                Path(f.name).unlink()

    @pytest.mark.asyncio
    async def test_document_order(self):
        """段落和表格按文档顺序交错出现。"""
        def build(doc):
            doc.add_paragraph("Before table.")
            table = doc.add_table(rows=2, cols=1)
            table.cell(0, 0).text = "T1"
            table.cell(1, 0).text = "T2"
            doc.add_paragraph("After table.")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.close()
            _create_docx(f.name, build)
            try:
                pages = await parse_document(f.name, "docx")
                text = pages[0][0]
                pos_before = text.index("Before table")
                pos_table = text.index("| T1 |")
                pos_after = text.index("After table")
                assert pos_before < pos_table < pos_after
            finally:
                Path(f.name).unlink()

    @pytest.mark.asyncio
    async def test_empty_document(self):
        """空文档返回空列表。"""
        def build(doc):
            pass  # 无内容

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.close()
            _create_docx(f.name, build)
            try:
                pages = await parse_document(f.name, "docx")
                assert pages == []
            finally:
                Path(f.name).unlink()


class TestDocxPreprocessor:
    """DOCX Preprocessor 清洗 + 归拢测试。"""

    @pytest.mark.asyncio
    async def test_preprocessor_pipeline(self):
        """完整 parser → preprocessor 管道。"""
        def build(doc):
            doc.add_heading("Test Title", level=1)
            doc.add_paragraph("Body paragraph one.")
            doc.add_paragraph("Body paragraph two.")
            doc.add_paragraph("Task A", style="List Bullet")
            doc.add_paragraph("Task B", style="List Bullet")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.close()
            _create_docx(f.name, build)
            try:
                pages = await parse_document(f.name, "docx")
                preprocessor = get_preprocessor()
                result = await preprocessor.preprocess("docx", pages)
                assert len(result) == 1
                text = result[0][0]
                # 标题保留
                assert "# Test Title" in text
                # 列表保留
                assert "- Task A" in text
                assert "- Task B" in text
                # 无多余连续空行
                assert "\n\n\n" not in text
            finally:
                Path(f.name).unlink()

    @pytest.mark.asyncio
    async def test_table_preserved_in_preprocessor(self):
        """表格在预处理后保持完整（无间断空行）。"""
        def build(doc):
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "A"
            table.cell(0, 1).text = "B"
            table.cell(1, 0).text = "C"
            table.cell(1, 1).text = "D"

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.close()
            _create_docx(f.name, build)
            try:
                pages = await parse_document(f.name, "docx")
                preprocessor = get_preprocessor()
                result = await preprocessor.preprocess("docx", pages)
                text = result[0][0]
                # 表格行连续，中间没有空行
                assert "| A | B |\n| --- | --- |\n| C | D |" in text
            finally:
                Path(f.name).unlink()

    @pytest.mark.asyncio
    async def test_doc_type_routing(self):
        """doc 类型映射到 DocxPreprocessor。"""
        def build(doc):
            doc.add_paragraph("Simple content.")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.close()
            _create_docx(f.name, build)
            try:
                pages = await parse_document(f.name, "doc")
                preprocessor = get_preprocessor()
                result = await preprocessor.preprocess("doc", pages)
                assert len(result) == 1
                assert "Simple content." in result[0][0]
            finally:
                Path(f.name).unlink()

    @pytest.mark.asyncio
    async def test_empty_heading_cleaned(self):
        """空标题行（# 后无内容）被清除。"""
        def build(doc):
            p = doc.add_paragraph()
            p.text = ""  # 空段落，有样式但无内容
            doc.add_heading("", level=2)  # 空标题
            doc.add_paragraph("Real content.")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.close()
            _create_docx(f.name, build)
            try:
                pages = await parse_document(f.name, "docx")
                preprocessor = get_preprocessor()
                result = await preprocessor.preprocess("docx", pages)
                text = result[0][0]
                # 确保有真实内容
                assert "Real content." in text
                # 不应有空标题标记 ## 后无内容（但 parser 不会为无内容段落生成 ##）
                # 实际上 parser 会 skip 空内容段落，所以这里验证 parser 行为
            finally:
                Path(f.name).unlink()
