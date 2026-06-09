import re
from collections.abc import Callable
from pathlib import Path

from loguru import logger


async def parse_document(file_path: str, doc_type: str) -> list[tuple[str, int | None]]:
    """解析文档内容，返回 [(text, page_number), ...]。
    对于非 PDF 文档，page_number 为 None。
    """
    path = Path(file_path)
    if not path.exists():
        logger.error(f"文件不存在: {file_path}")
        return []

    parser = _PARSERS.get(doc_type)
    if parser is None:
        logger.warning(f"不支持的文档类型: {doc_type}，跳过解析")
        return []

    try:
        pages = await parser(path)
        if not pages:
            logger.warning(f"文档解析结果为空: {file_path}")
            return []
        return pages
    except Exception as exc:
        logger.error(f"文档解析失败: {file_path} — {exc}")
        return []


async def _parse_txt(path: Path) -> list[tuple[str, int | None]]:
    text = path.read_text(encoding="utf-8", errors="replace").strip()
    if not text:
        return []
    return [(text, None)]


async def _parse_md(path: Path) -> list[tuple[str, int | None]]:
    text = path.read_text(encoding="utf-8", errors="replace").strip()
    if not text:
        return []
    return [(text, None)]


async def _parse_pdf(path: Path) -> list[tuple[str, int | None]]:
    pages = await _parse_pdf_pypdf2(path)
    if pages:
        return pages
    # PyPDF2 失败，回退到 pdfplumber
    return await _parse_pdf_plumber(path)


async def _parse_pdf_pypdf2(path: Path) -> list[tuple[str, int | None]]:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        logger.warning("PyPDF2 未安装")
        return []

    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        logger.warning(f"PyPDF2 打开 PDF 失败: {exc}")
        return []

    pages: list[tuple[str, int | None]] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text()
        except Exception:
            continue
        if text and text.strip():
            pages.append((text.strip(), i + 1))  # 1-based 页码
    return pages


async def _parse_pdf_plumber(path: Path) -> list[tuple[str, int | None]]:
    try:
        import pdfplumber
    except ImportError:
        logger.error("pdfplumber 未安装，无法回退解析 PDF")
        return []

    pages: list[tuple[str, int | None]] = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages.append((text.strip(), i + 1))
    except Exception as exc:
        logger.error(f"pdfplumber 解析 PDF 失败: {exc}")
        return []
    return pages


async def _parse_docx(path: Path) -> list[tuple[str, int | None]]:
    """解析 DOCX 文档，保留结构和样式信息。

    按文档顺序处理段落和表格：
      - Heading 1-6 样式 → # ~ ###### 标题前缀
      - 粗体 + 大字号（≥14pt）段落 → ## 标题候选
      - 列表项（含 numPr 编号属性）→ - 前缀
      - 表格 → markdown 表格格式（| col1 | col2 |）
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        logger.error("python-docx 未安装")
        return []

    try:
        doc = Document(str(path))
    except Exception as exc:
        logger.error(f"docx_open_failed path={path} error={exc}")
        return []

    output_parts: list[str] = []

    # 按文档顺序遍历 body 子元素（段落 + 表格交错出现）
    body = doc.element.body
    para_idx = 0
    table_idx = 0
    paragraphs = doc.paragraphs
    tables = doc.tables

    _P_TAG = qn("w:p")
    _TBL_TAG = qn("w:tbl")

    for child in body:
        tag = child.tag

        if tag == _P_TAG:
            if para_idx < len(paragraphs):
                formatted = _format_docx_paragraph(paragraphs[para_idx])
                if formatted:
                    output_parts.append(formatted)
                para_idx += 1

        elif tag == _TBL_TAG:
            if table_idx < len(tables):
                table_md = _format_docx_table(tables[table_idx])
                if table_md:
                    output_parts.append(table_md)
                table_idx += 1

    # 兜底：处理 body 遍历未覆盖的段落/表格（理论上不会发生）
    while para_idx < len(paragraphs):
        formatted = _format_docx_paragraph(paragraphs[para_idx])
        if formatted:
            output_parts.append(formatted)
        para_idx += 1
    while table_idx < len(tables):
        table_md = _format_docx_table(tables[table_idx])
        if table_md:
            output_parts.append(table_md)
        table_idx += 1

    if not output_parts:
        return []
    return [("\n\n".join(output_parts), None)]


# ═══════════════════════════════════════════════════════════════
# DOCX 格式转换辅助函数
# ═══════════════════════════════════════════════════════════════

def _format_docx_paragraph(para) -> str:
    """将 DOCX 段落转换为带格式标记的文本。

    检测顺序：
      1. 样式名包含 Heading → # 标题
      2. 粗体 + 字号 ≥ 14pt → ## 标题候选
      3. 编号属性 numPr → - 列表项
      4. 其他 → 原样返回
    """
    text = para.text
    if text is None or not text.strip():
        return ""

    text = text.strip()

    # ① 样式检测
    style_name = ""
    if para.style and para.style.name:
        style_name = para.style.name

    heading_level = _detect_heading_level(style_name, para)
    if heading_level:
        prefix = "#" * heading_level
        return f"{prefix} {text}"

    # ② 列表检测（含 numPr 编号属性或 bullet 样式）
    if _is_list_paragraph(para):
        return f"- {text}"

    return text


def _detect_heading_level(style_name: str, para) -> int:
    """检测段落标题级别。

    Returns:
        1-6 表示标题级别，0 表示非标题。
    """
    # 标准 Heading 样式
    _HEADING_STYLE_RE = re.compile(r"^heading\s*(\d)", re.IGNORECASE)
    match = _HEADING_STYLE_RE.match(style_name)
    if match:
        level = int(match.group(1))
        return min(level, 6)

    # Title / Subtitle 样式
    if style_name.lower() in ("title",):
        return 1
    if style_name.lower() in ("subtitle",):
        return 2

    # TOC 目录样式 → 根据级别
    _TOC_RE = re.compile(r"^TOC\s*(\d)", re.IGNORECASE)
    match = _TOC_RE.match(style_name)
    if match:
        level = int(match.group(1))
        return min(level, 6)

    # 启发式：粗体 + 大字号 → 标题候选
    if para.runs:
        has_bold = False
        max_size = 0
        for run in para.runs:
            if run.bold:
                has_bold = True
            if run.font.size and run.font.size.emu:
                max_size = max(max_size, run.font.size.emu)

        # 14pt = 177800 EMU, 16pt = 203200 EMU
        if has_bold and max_size >= 177800:
            return 2  # 二级标题

    return 0


def _is_list_paragraph(para) -> bool:
    """检测段落是否有编号/列表属性。

    DOCX 列表标识可能位于两处：
      1. 段落自身的 w:pPr > w:numPr（直接编号）
      2. 所引用样式的 w:pPr > w:numPr（样式编号，如 "List Number" 样式）
    """
    from docx.oxml.ns import qn

    # ① 检测段落自身的 numPr
    pPr = para._element.find(qn("w:pPr"))
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            return True

    # ② 检测样式的 numPr + 样式名
    style = para.style
    if style is not None:
        style_name = style.name or ""
        if "list" in style_name.lower():
            return True
        # 也可以检测样式中是否定义了 numPr
        style_pPr = style._element.find(qn("w:pPr"))
        if style_pPr is not None:
            style_numPr = style_pPr.find(qn("w:numPr"))
            if style_numPr is not None:
                return True

    return False


def _format_docx_table(table) -> str:
    """将 DOCX 表格转换为 markdown 表格格式。

    第一行视为表头，自动插入分隔行。
    跳过全空单元格的行（可能是装饰性边框）。
    """
    if not table.rows:
        return ""

    rows: list[str] = []
    col_count = len(table.rows[0].cells) if table.rows else 0

    if col_count == 0:
        return ""

    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        # 跳过全空行
        if all(not c for c in cells):
            continue
        # 补齐列数（合并单元格可能导致列数不一致）
        while len(cells) < col_count:
            cells.append("")
        rows.append("| " + " | ".join(cells) + " |")

    if not rows:
        return ""

    # 插入表头分隔行（第一行之后）
    if len(rows) >= 1:
        sep = "| " + " | ".join(["---"] * col_count) + " |"
        rows.insert(1, sep)

    return "\n".join(rows)


_PARSERS: dict[str, Callable[..., object]] = {
    "txt": _parse_txt,
    "md": _parse_md,
    "pdf": _parse_pdf,
    "docx": _parse_docx,
    "doc": _parse_docx,
}
